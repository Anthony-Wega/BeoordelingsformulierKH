import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import zipfile
from datetime import date
import os
import re

# Page config
st.set_page_config(
    page_title="Eindbeoordeling Generator",
    page_icon="📝",
    layout="wide"
)

st.title("📝 Automatische Eindbeoordeling Generator")
st.markdown("Genereer automatisch eindbeoordeling formulieren voor KH1 AGZ PvB")

# Template path
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "BF KH1 AGZ PvB Eindbeoordeling.docx")


def format_score(value):
    """Format score as Dutch decimal (comma instead of period)"""
    if pd.isna(value):
        return ""
    return str(round(float(value), 1)).replace(".", ",")


def load_excel(uploaded_file):
    """Load and parse the Excel file"""
    df = pd.read_excel(uploaded_file, header=None)

    # Find the header row (contains 'Studentnummer')
    header_row = None
    for i, row in df.iterrows():
        if 'Studentnummer' in row.values:
            header_row = i
            break

    if header_row is None:
        st.error("Kon geen header rij vinden met 'Studentnummer'")
        return None

    # Set headers and filter data
    df.columns = df.iloc[header_row]
    df = df.iloc[header_row + 1:]

    # Remove empty rows
    df = df.dropna(subset=['Studentnummer'])
    df = df[df['Studentnummer'].notna()]

    # Select relevant columns
    required_cols = ['Studentnummer', 'Naam', 'Anam/LO', 'Verslag', 'Reflectie',
                     'Onderdeel A', 'Onderdeel B', 'Onderdeel C', 'Totaal', 'Eindscore']

    available_cols = [col for col in required_cols if col in df.columns]
    df = df[available_cols].copy()

    # Convert studentnummer to int for display
    df['Studentnummer'] = df['Studentnummer'].astype(int)

    return df


def fill_template(template_path, student_data, common_data):
    """Fill the Word template with student and common data"""
    doc = Document(template_path)

    # Table 1: Basic info
    table1 = doc.tables[0]

    # Row 0: Naam VioS (student name with number)
    naam = student_data['Naam']
    studentnummer = int(student_data['Studentnummer'])
    table1.rows[0].cells[1].text = f"{naam} ({studentnummer})"

    # Row 1: Naam examinator
    table1.rows[1].cells[1].text = common_data['examinator']

    # Row 2: Datum beoordeling
    table1.rows[2].cells[1].text = common_data['datum']

    # Row 3: Gelegenheid
    table1.rows[3].cells[1].text = common_data['gelegenheid']

    # Row 4: Score (eindscore)
    table1.rows[4].cells[1].text = format_score(student_data['Eindscore'])

    # Table 2: Score overview
    table2 = doc.tables[1]

    # Row 2: Anamnese (update the cell text preserving structure)
    anam_score = format_score(student_data['Anam/LO'])
    onderdeel_a = format_score(student_data['Onderdeel A'])

    # Get original text and update score
    cell_text = table2.rows[2].cells[0].text
    # Replace the score in the text (looking for pattern like "5,7" at the end)
    new_text = re.sub(r'\d+[,\.]\d+$', anam_score, cell_text.strip())
    if not re.search(r'\d+[,\.]\d+$', cell_text.strip()):
        new_text = cell_text.rstrip() + "  " + anam_score
    table2.rows[2].cells[0].text = new_text
    table2.rows[2].cells[2].text = onderdeel_a

    # Row 3: Verslag
    verslag_score = format_score(student_data['Verslag'])
    onderdeel_b = format_score(student_data['Onderdeel B'])

    cell_text = table2.rows[3].cells[0].text
    new_text = re.sub(r'\d+[,\.]\d+$', verslag_score, cell_text.strip())
    if not re.search(r'\d+[,\.]\d+$', cell_text.strip()):
        new_text = cell_text.rstrip() + "  " + verslag_score
    table2.rows[3].cells[0].text = new_text
    table2.rows[3].cells[2].text = onderdeel_b

    # Row 4: Reflectie
    reflectie_score = format_score(student_data['Reflectie'])
    onderdeel_c = format_score(student_data['Onderdeel C'])

    cell_text = table2.rows[4].cells[0].text
    new_text = re.sub(r'\d+[,\.]\d+$', reflectie_score, cell_text.strip())
    if not re.search(r'\d+[,\.]\d+$', cell_text.strip()):
        new_text = cell_text.rstrip() + "  " + reflectie_score
    table2.rows[4].cells[0].text = new_text
    table2.rows[4].cells[2].text = onderdeel_c

    # Row 5: Totaal
    totaal = format_score(student_data['Totaal'])
    table2.rows[5].cells[2].text = totaal

    return doc


def generate_all_forms(df, template_path, common_data):
    """Generate all forms and return as ZIP file"""
    zip_buffer = BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for idx, row in df.iterrows():
            student_data = row.to_dict()
            doc = fill_template(template_path, student_data, common_data)

            # Save document to BytesIO
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Create filename from student name
            naam = student_data['Naam'].replace(',', '').replace(' ', '_')
            studentnummer = int(student_data['Studentnummer'])
            filename = f"Eindbeoordeling_{naam}_{studentnummer}.docx"

            zip_file.writestr(filename, doc_buffer.getvalue())

    zip_buffer.seek(0)
    return zip_buffer


# Main app
st.header("1. Gemeenschappelijke Gegevens")
st.markdown("Deze gegevens worden op alle formulieren ingevuld.")

col1, col2, col3 = st.columns(3)

with col1:
    examinator = st.text_input("Naam examinator", placeholder="Bijv. Jan de Vries")

with col2:
    datum = st.date_input("Datum beoordeling", value=date.today())
    datum_str = datum.strftime("%d-%m-%Y")

with col3:
    gelegenheid = st.selectbox("Gelegenheid", ["1e", "Herkansing"])

st.header("2. Excel Bestand Uploaden")

uploaded_file = st.file_uploader(
    "Upload het Excel bestand met studentgegevens",
    type=['xlsx', 'xls'],
    help="Het bestand moet kolommen hebben voor: Studentnummer, Naam, Anam/LO, Verslag, Reflectie, Onderdeel A/B/C, Totaal, Eindscore"
)

if uploaded_file is not None:
    df = load_excel(uploaded_file)

    if df is not None:
        st.header("3. Preview Studenten")
        st.markdown(f"**{len(df)} studenten gevonden**")

        # Show preview
        display_df = df.copy()
        display_df['Eindscore'] = display_df['Eindscore'].apply(lambda x: format_score(x))
        st.dataframe(display_df, use_container_width=True)

        st.header("4. Formulieren Genereren")

        if not examinator:
            st.warning("Vul eerst de naam van de examinator in.")
        else:
            common_data = {
                'examinator': examinator,
                'datum': datum_str,
                'gelegenheid': gelegenheid
            }

            if st.button("🚀 Genereer Formulieren", type="primary"):
                with st.spinner("Formulieren worden gegenereerd..."):
                    try:
                        zip_file = generate_all_forms(df, TEMPLATE_PATH, common_data)

                        st.success(f"✅ {len(df)} formulieren succesvol gegenereerd!")

                        st.download_button(
                            label="📥 Download ZIP met alle formulieren",
                            data=zip_file,
                            file_name=f"Eindbeoordelingen_{datum_str}.zip",
                            mime="application/zip"
                        )
                    except Exception as e:
                        st.error(f"Er ging iets mis: {str(e)}")
                        st.exception(e)

else:
    st.info("Upload een Excel bestand om te beginnen.")

# Footer
st.markdown("---")
st.markdown("*Eindbeoordeling Generator voor KH1 AGZ PvB*")
